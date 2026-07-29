"""
extractor.py - Document text extraction module for PDF and DOCX formats.
Handles file validation, multi-page parsing, layout preservation, and error handling.
"""

import os
import io
from typing import Union, BinaryIO
import fitz  # PyMuPDF
import docx

from utils import logger, preprocess_text


class ResumeExtractorError(Exception):
    """Base exception class for resume extraction errors."""
    pass


class UnsupportedFileFormatError(ResumeExtractorError):
    """Raised when an unsupported file format is provided."""
    pass


class CorruptedFileError(ResumeExtractorError):
    """Raised when a PDF or DOCX document cannot be parsed due to corruption."""
    pass


class EmptyFileError(ResumeExtractorError):
    """Raised when the document contains no readable text."""
    pass


class ResumeExtractor:
    """Class to manage document text extraction from PDF and DOCX files."""

    @staticmethod
    def extract_from_pdf(file_source: Union[str, bytes, BinaryIO]) -> str:
        """Extracts plain text from a PDF file using PyMuPDF (fitz).

        Args:
            file_source: File path, bytes object, or file-like object.

        Returns:
            str: Extracted and cleaned text from PDF.

        Raises:
            CorruptedFileError: If PyMuPDF fails to open or parse document.
            EmptyFileError: If extracted text is empty.
        """
        try:
            if isinstance(file_source, (str, os.PathLike)):
                doc = fitz.open(file_source)
            elif isinstance(file_source, bytes):
                doc = fitz.open(stream=file_source, filetype="pdf")
            else:
                # File-like object (e.g. Streamlit UploadedFile)
                content = file_source.read()
                if hasattr(file_source, 'seek'):
                    file_source.seek(0)
                doc = fitz.open(stream=content, filetype="pdf")

            full_text = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                # 'text' mode preserves reading layout order
                page_text = page.get_text("text")
                if page_text:
                    full_text.append(page_text)

            doc.close()
            combined_text = "\n".join(full_text)

            cleaned = preprocess_text(combined_text)
            if not cleaned:
                raise EmptyFileError("The uploaded PDF file contains no readable text or is scanned/image-only.")

            return cleaned

        except EmptyFileError:
            raise
        except Exception as e:
            logger.error(f"Failed to extract PDF text: {str(e)}")
            raise CorruptedFileError(f"Could not parse PDF file. The file may be damaged or corrupted: {str(e)}")

    @staticmethod
    def extract_from_docx(file_source: Union[str, bytes, BinaryIO]) -> str:
        """Extracts plain text from a DOCX file using python-docx.

        Args:
            file_source: File path, bytes object, or file-like object.

        Returns:
            str: Extracted and cleaned text from DOCX.

        Raises:
            CorruptedFileError: If python-docx fails to open or parse document.
            EmptyFileError: If extracted text is empty.
        """
        try:
            if isinstance(file_source, (str, os.PathLike)):
                doc = docx.Document(file_source)
            elif isinstance(file_source, bytes):
                doc = docx.Document(io.BytesIO(file_source))
            else:
                content = file_source.read()
                if hasattr(file_source, 'seek'):
                    file_source.seek(0)
                doc = docx.Document(io.BytesIO(content))

            full_text = []

            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    full_text.append(paragraph.text)

            # Extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            combined_text = "\n".join(full_text)
            cleaned = preprocess_text(combined_text)

            if not cleaned:
                raise EmptyFileError("The uploaded DOCX file contains no readable text.")

            return cleaned

        except EmptyFileError:
            raise
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {str(e)}")
            raise CorruptedFileError(f"Could not parse DOCX file. The file may be damaged or corrupted: {str(e)}")

    @classmethod
    def extract_text(cls, file_source: Union[str, bytes, BinaryIO], filename: str = "") -> str:
        """Main interface to extract text based on filename or file extension.

        Args:
            file_source: File path or stream.
            filename: File name string used to determine format extension.

        Returns:
            str: Cleaned text from document.

        Raises:
            UnsupportedFileFormatError: If file type is not .pdf or .docx.
        """
        if isinstance(file_source, (str, os.PathLike)):
            ext = os.path.splitext(file_source)[1].lower()
        else:
            ext = os.path.splitext(filename)[1].lower() if filename else ""

        if ext == ".pdf":
            return cls.extract_from_pdf(file_source)
        elif ext == ".docx":
            return cls.extract_from_docx(file_source)
        else:
            raise UnsupportedFileFormatError(
                f"Unsupported file extension '{ext}'. Only .pdf and .docx files are supported."
            )
